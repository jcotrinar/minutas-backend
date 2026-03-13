"""
generador_minutas.py
Reemplaza los marcadores «VARIABLE» en los templates Word.
Equivalente al Módulo1 (contratos) del VBA original.
"""

import os
import re
import copy
from pathlib import Path
from datetime import date
from docx import Document
from docx.oxml.ns import qn
from lxml import etree
from app.services.calculos import compilar_variables

TEMPLATES_DIR = Path(os.getenv("TEMPLATES_DIR", "app/templates"))
OUTPUT_DIR    = Path(os.getenv("OUTPUT_DIR", "minutas_generadas"))

COLORES_TEMPLATE = {
    "VERDE":    "VERDE.docx",
    "AMARILLO": "AMARILLO.docx",
    "AZUL":     "AZUL.docx",
    "ROJO":     "ROJO.docx",
}

# Normaliza la clave del marcador: quita espacios internos y strips
def _norm(clave: str) -> str:
    return clave.strip()

def _build_lookup(variables: dict) -> dict:
    """Crea un dict con claves normalizadas para tolerar espacios en los marcadores."""
    return {_norm(k): v for k, v in variables.items()}


def _reemplazar_en_parrafo(parrafo, lookup: dict):
    """
    Reemplaza marcadores «VARIABLE» preservando el formato (negrita, tamaño, etc.)
    de cada run individual.
    
    Estrategia:
    1. Reconstruir el texto completo del párrafo detectando posición de cada run.
    2. Por cada marcador encontrado, reemplazar directamente en el run que lo contiene
       (o en el primer run del span si está partido entre varios).
    3. Preservar rPr (formato) de todos los runs.
    """
    # Texto completo
    texto = "".join(r.text for r in parrafo.runs)
    patron = re.compile(r"«([^»]+)»")
    
    if not patron.search(texto):
        return  # Sin marcadores

    # Reconstruir texto run por run con sus posiciones
    runs = parrafo.runs
    if not runs:
        return

    # Índice de inicio de cada run dentro del texto completo
    posiciones = []
    pos = 0
    for r in runs:
        posiciones.append(pos)
        pos += len(r.text)

    # Para cada marcador, encontrar en qué runs cae y reemplazar
    # Procesamos de atrás hacia adelante para no desplazar índices
    matches = list(patron.finditer(texto))
    for match in reversed(matches):
        clave = _norm(match.group(1))
        valor = lookup.get(clave, match.group(0))
        start, end = match.start(), match.end()

        # Encontrar runs que cubren [start, end)
        runs_span = []
        for i, r in enumerate(runs):
            r_start = posiciones[i]
            r_end   = r_start + len(r.text)
            if r_end > start and r_start < end:
                runs_span.append(i)

        if not runs_span:
            continue

        first_i = runs_span[0]
        last_i  = runs_span[-1]

        # Calcular texto antes y después del marcador dentro del primer y último run
        r_first = runs[first_i]
        r_last  = runs[last_i]

        before = r_first.text[: start - posiciones[first_i]]
        after  = r_last.text[end - posiciones[last_i] :]

        # Poner valor en el primer run
        r_first.text = before + valor + after

        # Vaciar runs intermedios y el último si es distinto al primero
        for i in runs_span[1:]:
            runs[i].text = ""

        # Actualizar posiciones (simplificado: recalcular desde cero)
        texto = "".join(r.text for r in runs)
        pos = 0
        for i, r in enumerate(runs):
            posiciones[i] = pos
            pos += len(r.text)


def _reemplazar_en_tabla(tabla, lookup: dict):
    for fila in tabla.rows:
        for celda in fila.cells:
            for parrafo in celda.paragraphs:
                _reemplazar_en_parrafo(parrafo, lookup)
            for subtabla in celda.tables:
                _reemplazar_en_tabla(subtabla, lookup)


def generar_minuta(contrato, lote, distrito1=None, distrito2=None) -> Path:
    estado = contrato.estado
    template_nombre = COLORES_TEMPLATE.get(estado)

    if not template_nombre:
        raise ValueError(f"Estado desconocido: {estado}")

    template_path = TEMPLATES_DIR / template_nombre
    if not template_path.exists():
        raise FileNotFoundError(
            f"Template no encontrado: {template_path}\n"
            f"Asegúrate de copiar los 4 archivos .docx a: {TEMPLATES_DIR}"
        )

    variables = compilar_variables(contrato, lote, distrito1, distrito2)
    lookup    = _build_lookup(variables)

    doc = Document(str(template_path))

    for parrafo in doc.paragraphs:
        _reemplazar_en_parrafo(parrafo, lookup)

    for tabla in doc.tables:
        _reemplazar_en_tabla(tabla, lookup)

    for seccion in doc.sections:
        for parrafo in seccion.header.paragraphs:
            _reemplazar_en_parrafo(parrafo, lookup)
        for parrafo in seccion.footer.paragraphs:
            _reemplazar_en_parrafo(parrafo, lookup)

    fecha       = contrato.fecha
    año         = str(fecha.year)
    mes         = f"{fecha.month:02d}"
    nombre_arch = f"MINUTA_{contrato.numero:04d}_{estado}_{fecha.strftime('%Y%m%d')}.docx"

    carpeta_salida = OUTPUT_DIR / año / mes
    carpeta_salida.mkdir(parents=True, exist_ok=True)

    ruta_salida = carpeta_salida / nombre_arch
    doc.save(str(ruta_salida))

    return ruta_salida